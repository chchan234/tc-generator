"""
텍스트 추출 모듈 - PDF, DOCX 파일에서 텍스트 추출 및 전처리
"""
import re
import PyPDF2
from docx import Document


class TextExtractor:
    """텍스트 추출 및 전처리 클래스"""
    
    def __init__(self):
        self.raw_text = ""
        self.processed_text = ""
    
    def extract_from_pdf(self, pdf_path):
        """
        PDF 파일에서 텍스트 추출
        
        Args:
            pdf_path (str): PDF 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                for page_num in range(num_pages):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            self.raw_text = text
            return text
        except Exception as e:
            print(f"PDF 텍스트 추출 오류: {str(e)}")
            return ""
    
    def extract_from_docx(self, docx_path):
        """
        DOCX 파일에서 텍스트 추출
        
        Args:
            docx_path (str): DOCX 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        try:
            doc = Document(docx_path)
            text = ""
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # 표(테이블) 내용 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            self.raw_text = text
            return text
        except Exception as e:
            print(f"DOCX 텍스트 추출 오류: {str(e)}")
            return ""
    
    def extract_text(self, file_path):
        """
        파일 확장자에 따라 적절한 텍스트 추출 함수 호출
        
        Args:
            file_path (str): 파일 경로
            
        Returns:
            str: 추출된 텍스트
        """
        if file_path.lower().endswith('.pdf'):
            return self.extract_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            return self.extract_from_docx(file_path)
        else:
            print(f"지원되지 않는 파일 형식: {file_path}")
            return ""
    
    def preprocess_text(self, text=None):
        """
        텍스트 전처리 (필터링, 불필요한 내용 제거)
        
        Args:
            text (str, optional): 전처리할 텍스트. None인 경우 self.raw_text 사용
            
        Returns:
            str: 전처리된 텍스트
        """
        if text is None:
            text = self.raw_text
        
        if not text:
            return ""
        
        # 여러 공백을 하나로 치환
        processed = re.sub(r'\s+', ' ', text)
        
        # 특수 문자 및 불필요한 기호 제거 (필요에 따라 조정)
        processed = re.sub(r'[^\w\s\.\,\:\;\(\)\[\]\-\'\"]', '', processed)
        
        # 빈 줄 제거
        processed = re.sub(r'\n\s*\n', '\n', processed)
        
        self.processed_text = processed
        return processed